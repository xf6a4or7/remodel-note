"""분석 JSON → 회의록 Word 문서(.docx).

이전 회의록과 동일한 6섹션 구성:
  1. 미팅 개요 / 2. 주요 결정사항(표) / 3. 미결·현장확인(표) /
  4. 고객 요청·선호 / 5. Action Items(표) / 6. 비고(견적 메모)
맑은 고딕 폰트, 네이비/블루 헤더 테마.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
HEADER_FILL = "2E75B6"
KOREAN_FONT = "맑은 고딕"


def _set_korean_font(run, size=None, bold=False, color=None):
    run.font.name = KOREAN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_korean_font(run, size=15 if level == 1 else 13, bold=True,
                     color=NAVY if level == 1 else BLUE)
    p.space_after = Pt(6)
    return p


def _table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr[i], HEADER_FILL)
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        _set_korean_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val) if val else "-")
            _set_korean_font(run, size=10)
    return table


def build_minutes(analysis: dict, out_path: str):
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    style.font.size = Pt(10.5)

    ov = analysis.get("meeting_overview", {})

    # 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title_p.add_run(ov.get("title", "상담 회의록"))
    _set_korean_font(trun, size=20, bold=True, color=NAVY)

    # 1. 개요
    _heading(doc, "1. 미팅 개요")
    if ov.get("date"):
        p = doc.add_paragraph()
        _set_korean_font(p.add_run(f"일시: {ov['date']}"), size=10.5)
    if ov.get("participants"):
        p = doc.add_paragraph()
        _set_korean_font(p.add_run("참석: " + ", ".join(ov["participants"])), size=10.5)
    if ov.get("summary"):
        p = doc.add_paragraph()
        _set_korean_font(p.add_run(ov["summary"]), size=10.5)

    # 2. 결정사항
    _heading(doc, "2. 주요 결정사항")
    decisions = analysis.get("decisions", [])
    _table(doc, ["항목", "결정 내용", "비고"],
           [[d.get("topic"), d.get("decision"), d.get("note")] for d in decisions]
           or [["-", "-", "-"]])

    # 3. 미결·현장확인
    _heading(doc, "3. 미결 · 현장확인 항목")
    pending = analysis.get("pending_items", [])
    pri_mark = {"high": "★ 높음", "medium": "중간", "low": "낮음"}
    _table(doc, ["항목", "우선순위", "필요 조치"],
           [[x.get("item"), pri_mark.get(x.get("priority", ""), x.get("priority", "")),
             x.get("action_needed")] for x in pending]
           or [["-", "-", "-"]])

    # 4. 고객 요청·선호
    _heading(doc, "4. 고객 요청 · 선호")
    reqs = analysis.get("customer_requirements", [])
    _table(doc, ["공간", "분류", "요청 내용"],
           [[r.get("space"), r.get("category"), r.get("requirement")] for r in reqs]
           or [["-", "-", "-"]])

    # 5. Action Items
    _heading(doc, "5. Action Items")
    actions = analysis.get("action_items", [])
    _table(doc, ["할 일", "담당", "기한"],
           [[a.get("task"), a.get("owner"), a.get("due")] for a in actions]
           or [["-", "-", "-"]])

    # 6. 비고(견적 메모)
    est = analysis.get("estimate_notes", [])
    if est:
        _heading(doc, "6. 비고 · 견적 메모")
        _table(doc, ["항목", "내용"],
               [[e.get("item"), e.get("detail")] for e in est])

    # python-docx 기본 settings.xml의 zoom 태그에 percent 속성 보강(검증 통과용)
    settings = doc.settings.element
    for z in settings.findall(qn("w:zoom")):
        if not z.get(qn("w:percent")):
            z.set(qn("w:percent"), "100")

    doc.save(out_path)
    return out_path
