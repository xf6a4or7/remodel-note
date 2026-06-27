"""회의록 docx 생성 테스트 — 유효한 Word 문서가 나오는지 검증."""
from docx import Document

from src.docgen import build_minutes

SAMPLE = {
    "meeting_overview": {
        "title": "테스트 상담",
        "date": "2026-06-27",
        "participants": ["강대표님", "여자고객님"],
        "summary": "주방 리모델링 1차 상담 요약.",
    },
    "decisions": [{"topic": "아일랜드 높이", "decision": "950mm 확정", "note": ""}],
    "pending_items": [{"item": "계단 위치", "priority": "high", "action_needed": "실측"}],
    "customer_requirements": [
        {"space": "주방", "category": "자재", "requirement": "BLUM 힌지", "raw_quote": ""}
    ],
    "action_items": [{"task": "견적 발송", "owner": "강대표님", "due": "이번 주"}],
    "estimate_notes": [{"item": "상판", "detail": "세닉스"}],
}


def test_build_minutes_creates_valid_docx(tmp_path):
    out = tmp_path / "회의록.docx"
    build_minutes(SAMPLE, str(out))
    assert out.exists() and out.stat().st_size > 0

    doc = Document(str(out))           # 열리면 유효한 docx
    assert len(doc.tables) == 5        # 결정/미결/요청/액션/견적 표
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "테스트 상담" in text


def test_build_minutes_handles_empty_sections(tmp_path):
    """빈 분석 결과여도 깨지지 않고 '-' 자리표시 표를 만든다."""
    out = tmp_path / "빈회의록.docx"
    build_minutes({"meeting_overview": {}}, str(out))
    doc = Document(str(out))
    assert len(doc.tables) >= 4        # 견적 메모는 비면 생략되므로 4개 이상
